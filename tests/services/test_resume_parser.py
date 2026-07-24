import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from weasyprint import HTML

from app.services.resume_parser import extract_text_from_docx, extract_text_from_pdf


def _add_docx_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_extract_text_from_docx_includes_hyperlink_urls():
    document = Document()
    document.add_paragraph("Vishesh Dudeja")
    paragraph = document.add_paragraph()
    _add_docx_hyperlink(paragraph, "https://github.com/vishesh0306", "GitHub")

    buf = io.BytesIO()
    document.save(buf)

    text = extract_text_from_docx(buf.getvalue())

    assert "Vishesh Dudeja" in text
    assert "https://github.com/vishesh0306" in text
    assert "'GitHub' -> https://github.com/vishesh0306" in text


def test_extract_text_from_pdf_includes_hyperlink_urls():
    html = (
        "<html><body><p>Vishesh Dudeja</p>"
        '<p>Find me on <a href="https://github.com/vishesh0306">GitHub</a></p>'
        "</body></html>"
    )
    pdf_bytes = HTML(string=html).write_pdf()

    text = extract_text_from_pdf(pdf_bytes)

    assert "Vishesh Dudeja" in text
    assert "https://github.com/vishesh0306" in text
