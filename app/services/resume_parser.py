import io

import pdfplumber
from docx import Document

from app.schemas.resume import ResumeContent
from app.services.llm_client import llm_client

RESUME_EXTRACTION_PROMPT = """You are extracting structured resume data from the raw resume text below.

Rules:
- Extract ONLY information explicitly present in the text.
- Do NOT invent, infer, guess, or embellish names, dates, companies, titles, skills, or achievements that are not literally stated in the text.
- If a field is not present in the source text, leave it as an empty string or an empty list.
- Preserve the original wording of experience bullet points as closely as possible; do not add quantification, metrics, or claims that are not already in the source.
- Every project often has several distinct bullet points — extract EACH one as its own entry in that project's `bullets` list. Do not compress multiple bullets into a single sentence and do not drop any of them.
- Skills are often grouped under labels beyond "languages/frameworks/tools/cloud": things like databases, operating systems, core CS concepts, or design patterns. Categorize into languages/frameworks/tools/cloud_devops where they clearly fit; put anything that doesn't fit one of those into `skills.other`. Never drop a stated skill for lack of a matching bucket.
- If the resume has a section for awards, hackathon wins, competition rankings, or other notable recognitions (often titled "Achievements", "Awards", "Honors"), extract each one as its own entry in `achievements`. This is distinct from `certifications` (credentials/courses) — do not mix the two, and do not drop this section.
- If the text below includes a "Hyperlinks found in document" list, it maps each link's visible label text to its actual URL. When you see a link label (e.g. "GitHub", "LinkedIn", a person's name used as a link, a project's "Video Demonstration" or "Project link" label) elsewhere in the resume, look up its real URL in that list and use the URL — never the bare label — in `contact.links` or a project's `link` field.

Resume text:
---
{text}
---
"""


class ResumeParsingError(Exception):
    """Raised when a resume file cannot be read or parsed into structured content."""


def _append_hyperlinks_block(text: str, link_lines: list[str]) -> str:
    if not link_lines:
        return text
    return text + "\n\n--- Hyperlinks found in document (label -> URL) ---\n" + "\n".join(link_lines)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            link_lines = []
            for page in pdf.pages:
                pages_text.append(page.extract_text() or "")
                for link in page.hyperlinks:
                    uri = link.get("uri")
                    if not uri:
                        continue
                    try:
                        label = (page.crop((link["x0"], link["top"], link["x1"], link["bottom"])).extract_text() or "").strip()
                    except Exception:
                        label = ""
                    label = label.replace("\n", " ")
                    link_lines.append(f"{label!r} -> {uri}" if label else uri)
    except Exception as exc:
        raise ResumeParsingError(f"Could not read PDF file: {exc}") from exc

    text = "\n".join(pages_text).strip()
    if not text:
        raise ResumeParsingError(
            "No extractable text found in this PDF (it may be a scanned image without a text layer)."
        )
    return _append_hyperlinks_block(text, link_lines)


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ResumeParsingError(f"Could not read DOCX file: {exc}") from exc

    text = "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
    if not text:
        raise ResumeParsingError("No extractable text found in this DOCX file.")

    link_lines = []
    for paragraph in document.paragraphs:
        for hyperlink in paragraph.hyperlinks:
            if not hyperlink.address:
                continue
            label = (hyperlink.text or "").strip()
            link_lines.append(f"{label!r} -> {hyperlink.address}" if label else hyperlink.address)

    return _append_hyperlinks_block(text, link_lines)


def parse_resume_text(text: str) -> ResumeContent:
    # Deliberately does not catch LLMError here — a failure at this point (quota exhaustion, an
    # unusable AI response) is an AI-service problem, not a problem with the candidate's file, and
    # ResumeParsingError (-> 400 Bad Request) would misreport it as one. Let it propagate to the
    # API layer's own LLMError handling, which distinguishes quota exhaustion (503) from other
    # failures (502).
    prompt = RESUME_EXTRACTION_PROMPT.format(text=text)
    return llm_client.generate_structured(prompt, ResumeContent)
